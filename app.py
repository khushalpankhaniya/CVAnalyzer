from flask import Flask, render_template, request, redirect, url_for
import os
import re
import pdfplumber
import docx

app = Flask(__name__)
UPLOAD_FOLDER = os.path.join("static", "uploads")
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def extract_text(filepath):
    ext = filepath.split(".")[-1].lower()
    text = ""
    num_pages = 1
    if ext == "pdf":
        with pdfplumber.open(filepath) as pdf:
            num_pages = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    elif ext in ["doc", "docx"]:
        doc = docx.Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        num_pages = len(doc.paragraphs) // 30 + 1
    return text, num_pages


KNOWN_SKILLS = [
    "Python", "Java", "C++", "C", "SQL", "HTML", "CSS", "JavaScript", "React",
    "Django", "Flask", "Node.js", "Express", "MongoDB", "AWS", "Docker",
    "Kubernetes", "Git", "REST", "Redux", "Bootstrap", "Tailwind", "TypeScript",
    "Data Analysis", "Machine Learning", "AI", "Firebase"
]


def extract_skills(text):
    skills_found = set()

   
    sections = re.split(r"(Skills|Technical Skills|Expertise|Competencies|Proficiencies)[:\n]", text, flags=re.I)
    if len(sections) > 1:
        possible_skills_text = sections[1]
        for skill in KNOWN_SKILLS:
            if re.search(r"\b" + re.escape(skill) + r"\b", possible_skills_text, re.I):
                skills_found.add(skill)

    for skill in KNOWN_SKILLS:
        if re.search(r"\b" + re.escape(skill) + r"\b", text, re.I):
            skills_found.add(skill)

    return ", ".join(sorted(skills_found)) if skills_found else "Not found"


def extract_education(text):

    sections = re.split(r"(Education|Qualifications|Academic Background|Certifications|Experience|Projects)[:\n]", text, flags=re.I)
    
    education_lines = []
    for i, sec in enumerate(sections):
        if re.search(r"Education|Qualifications|Academic Background", sec, re.I):
            if i + 1 < len(sections):
                lines = sections[i + 1].split("\n")
                for line in lines:
                    line = line.strip()
                    if line:
                      
                        if re.search(r"(Diploma|Bachelor|B\.Sc|M\.Sc|MBA|B\.Tech|M\.Tech|PhD|B\.Com|M\.Com|B\.A|M\.A|BBA|High School|SSC|HSC|BE)", line, re.I):
                            education_lines.append(line)
            break

    if not education_lines:
        education_lines = re.findall(r"(Diploma|Bachelor|B\.Sc|M\.Sc|MBA|B\.Tech|M\.Tech|PhD|B\.Com|M\.Com|B\.A|M\.A|BBA|High School|SSC|HSC|BE)[^,\n]*", text, re.I)

    return ", ".join([e.strip() for e in education_lines]) if education_lines else "Not found"


def parse_resume(filepath):
    text, num_pages = extract_text(filepath)
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # Name: first line with letters
    name = next((line for line in lines[:5] if re.match(r"^[A-Za-z\s]{2,}$", line)), "Unknown")

    # Email
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    email = email_match.group() if email_match else "Not found"

    # Phone
    phone_match = re.search(r"(\+?\d[\d\s\-\(\)]{7,}\d)", text)
    phone = phone_match.group() if phone_match else "Not found"

    # Skills
    skills = extract_skills(text)

    # Education
    education = extract_education(text)

    # Experience
    exp_match = re.findall(r"(\d+\+?\s+years)", text, re.I)
    experience = exp_match[0] if exp_match else "Not found"

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "education": education,
        "skills": skills,
        "experience": experience,
        "pages": num_pages,
        "filename": os.path.basename(filepath)
    }


# Routes
@app.route("/", methods=["GET", "POST"])
def upload_resume():
    if request.method == "POST":
        file = request.files.get("resume")
        if file and file.filename != "":
            filepath = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
            file.save(filepath)
            return redirect(url_for("dashboard"))
    return render_template("index.html")


@app.route("/dashboard")
def dashboard():
    files = os.listdir(app.config["UPLOAD_FOLDER"])
    resumes = [parse_resume(os.path.join(app.config["UPLOAD_FOLDER"], f)) for f in files]
    return render_template("dashboard.html", resumes=resumes)


if __name__ == "__main__":
    app.run(debug=True)
